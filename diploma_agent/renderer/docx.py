"""DOCX export for Word-based thesis editing."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from diploma_agent import config
from diploma_agent.state import Section, ThesisState, ensure_workspace, read_section


@dataclass(frozen=True)
class DocxResult:
    docx_path: Path
    message: str


def render_docx(state: ThesisState, workspace: Path | None = None) -> DocxResult:
    root = ensure_workspace(workspace or config.workspace_dir())
    build = root / "build"
    docx_path = build / "diploma.docx"

    doc = Document()
    _configure_document(doc)
    _add_title_page(doc, state)
    _add_toc_stub(doc, state)
    _add_sections(doc, state, root, seen_images=set())
    _add_bibliography(doc, root)
    doc.save(docx_path)

    return DocxResult(docx_path, "DOCX собран.")


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for idx in range(1, 4):
        style = doc.styles[f"Heading {idx}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(16 if idx == 1 else 14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def _add_title_page(doc: Document, state: ThesisState) -> None:
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ДИПЛОМНАЯ РАБОТА")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)

    topic = doc.add_paragraph()
    topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    topic.paragraph_format.first_line_indent = Cm(0)
    topic.add_run(f"\n{state.topic}").font.size = Pt(14)

    for _ in range(8):
        doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.first_line_indent = Cm(0)
    footer.add_run("Москва, 2026").font.size = Pt(14)
    doc.add_section(WD_SECTION.NEW_PAGE)


def _add_toc_stub(doc: Document, state: ThesisState) -> None:
    _add_centered_heading(doc, "СОДЕРЖАНИЕ", level=1)
    for section_id, section in sorted(state.sections.items(), key=_section_sort_key):
        if not section.file:
            continue
        indent = "    " if "." in section_id else ""
        p = doc.add_paragraph(f"{indent}{section.title}")
        p.paragraph_format.first_line_indent = Cm(0)
    doc.add_section(WD_SECTION.NEW_PAGE)


def _add_sections(doc: Document, state: ThesisState, workspace: Path, seen_images: set[str]) -> None:
    added_any = False
    for _, section in sorted(state.sections.items(), key=_section_sort_key):
        text = read_section(section, workspace)
        if not text.strip():
            continue
        if added_any and (section.id in {"intro", "1", "2", "3", "4", "5", "conclusion"}):
            doc.add_section(WD_SECTION.NEW_PAGE)
        _add_markdown(doc, text, workspace / "sections", seen_images)
        added_any = True


def _add_markdown(doc: Document, markdown: str, base_dir: Path, seen_images: set[str] | None = None) -> None:
    seen_images = seen_images if seen_images is not None else set()
    in_math = False
    math_lines: list[str] = []
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if _should_skip_line(line):
            continue
        if line.strip() == "$$":
            if in_math:
                _add_formula_block(doc, "\n".join(math_lines))
                math_lines = []
            in_math = not in_math
            continue
        if in_math:
            math_lines.append(line)
            continue
        if not line:
            continue

        image = re.match(r"^!\[(.*?)\]\((.*?)\)\s*$", line)
        if image:
            caption, path = image.groups()
            _add_image(doc, base_dir, path, caption, seen_images)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            _add_centered_heading(doc, _normalize_heading_text(heading.group(2)), level=level)
            continue

        bullet = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bullet:
            p = doc.add_paragraph(_strip_inline(bullet.group(1)), style="List Bullet")
            p.paragraph_format.first_line_indent = Cm(0)
            continue

        numbered = re.match(r"^\s*\d+[.)]\s+(.*)$", line)
        if numbered:
            p = doc.add_paragraph(_strip_inline(line))
            p.paragraph_format.first_line_indent = Cm(0)
            continue

        if _looks_unfinished(line):
            continue
        doc.add_paragraph(_strip_inline(line))
    if math_lines:
        _add_formula_block(doc, "\n".join(math_lines))


def _add_image(doc: Document, base_dir: Path, image_path: str, caption: str, seen_images: set[str]) -> None:
    path = (base_dir / image_path).resolve()
    key = str(path).lower()
    if key in seen_images:
        return
    seen_images.add(key)
    if not path.exists():
        missing = doc.add_paragraph(f"[Рисунок не найден: {image_path}]")
        missing.paragraph_format.first_line_indent = Cm(0)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(15))
    if caption:
        c = doc.add_paragraph(f"Рисунок - {_strip_inline(caption)}")
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.first_line_indent = Cm(0)


def _add_formula_block(doc: Document, formula: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(_pretty_formula(formula))
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def _add_bibliography(doc: Document, workspace: Path) -> None:
    path = workspace / "build" / "bibliography.md"
    if not path.exists():
        return
    doc.add_section(WD_SECTION.NEW_PAGE)
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            _add_centered_heading(doc, _normalize_heading_text(heading.group(2)), level=min(len(heading.group(1)), 3))
            continue
        p = doc.add_paragraph(_strip_inline(line))
        p.paragraph_format.first_line_indent = Cm(0)


def _strip_inline(text: str) -> str:
    text = _pretty_formula(text)
    text = text.replace(r"\(", "").replace(r"\)", "")
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\*\*([^*]*)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]*)\*", r"\1", text)
    text = text.replace("$", "")
    return text.strip()


def _add_centered_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(16 if level == 1 else 14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)


def _normalize_heading_text(text: str) -> str:
    clean = _strip_inline(text)
    special = {
        "введение": "ВВЕДЕНИЕ",
        "содержание": "СОДЕРЖАНИЕ",
        "заключение": "ЗАКЛЮЧЕНИЕ",
        "список использованных источников": "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
    }
    return special.get(clean.lower(), clean)


def _pretty_formula(text: str) -> str:
    replacements = {
        r"\frac{|C|}{|F|}": "|C| / |F|",
        r"\frac{coverage}{distance}": "coverage / distance",
        r"\frac{\sigma(d_1, d_2, ..., d_n)}{\mu(d_1, d_2, ..., d_n)}": "sigma(d1, d2, ..., dn) / mu(d1, d2, ..., dn)",
        r"\min_{r,p,t}": "min_{r,p,t}",
        r"\|": "||",
        r"\sigma": "sigma",
        r"\mu": "mu",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def _should_skip_line(line: str) -> bool:
    stripped = line.strip()
    forbidden = (
        "автоматически подготовленный",
        "Текст требует ручной",
        "Ключевые опорные материалы",
        "ТРЕБУЕТ УТОЧНЕНИЯ",
        "[ТРЕБУЕТ УТОЧНЕНИЯ]",
    )
    if any(item.lower() in stripped.lower() for item in forbidden):
        return True
    if stripped.startswith("- `docs/") or stripped.startswith("- `coverage_") or stripped.startswith("- `results/"):
        return True
    return False


def _looks_unfinished(line: str) -> bool:
    stripped = _strip_inline(line)
    if len(stripped) < 80:
        return False
    if stripped.endswith((".", "!", "?", "…", ":", ";", ")", "»")):
        return False
    return True


def _section_sort_key(item: tuple[str, Section]) -> tuple[int, ...]:
    section_id, _ = item
    special = {"annotation": (0,), "intro": (1,), "conclusion": (9998,)}
    if section_id in special:
        return special[section_id]
    parts = []
    for part in section_id.split("."):
        try:
            parts.append(int(part))
        except ValueError:
            parts.append(9999)
    return tuple(parts)
